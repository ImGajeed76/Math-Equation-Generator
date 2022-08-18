from docx import Document

from math_problem_creator.equation_generator import EquationCreator


class WordGenerator:
    table = []

    def create_table(self, num_rows: int, num_cols: int, equation_fun: EquationCreator):
        table = []
        for i in range(num_rows):
            row = []
            for j in range(num_cols):
                row.append(equation_fun.run())
            table.append(row)
        self.table = table
        return table

    def create_simple_word_doc2(self, table_style: str = "Table Grid"):
        document = Document()

        solve_space = document.add_table(rows=len(self.table), cols=len(self.table[0]) * 2)
        solve_space.style = table_style

        for i in range(len(solve_space.rows)):
            cells = solve_space.rows[i].cells
            for j in range(int(len(cells) / 2)):
                cells[j * 2].text = f"{j * len(self.table) + i + 1}"

        document.add_paragraph("")
        document.add_paragraph("Equations")

        equations = document.add_table(rows=len(self.table), cols=len(self.table[0]) * 2)
        equations.style = table_style

        for i in range(len(equations.rows)):
            cells = equations.rows[i].cells
            for j in range(int(len(cells) / 2)):
                cells[j * 2].text = f"{j * len(self.table) + i + 1}"
                cells[j * 2 + 1].text = f"{self.table[i][j].equation}"

        document.add_page_break()
        document.add_paragraph("Solutions")

        solutions = document.add_table(rows=len(self.table), cols=len(self.table[0]) * 2)
        solutions.style = table_style

        for i in range(len(solutions.rows)):
            cells = solutions.rows[i].cells
            for j in range(int(len(cells) / 2)):
                cells[j * 2].text = f"{j * len(self.table) + i + 1}"
                cells[j * 2 + 1].text = f"{self.table[i][j].solution}"

        document.save("math_problem_creator.docx")

    def create_simple_word_doc(self, table_style: str = "Table Grid"):
        document = Document()

        equations = document.add_table(rows=len(self.table), cols=len(self.table[0]) * 2)
        equations.style = table_style

        for i in range(len(equations.rows)):
            cells = equations.rows[i].cells
            for j in range(int(len(cells) / 2)):
                cells[j * 2].text = f"{self.table[i][j].equation}"

        document.add_page_break()
        document.add_paragraph("Solutions")

        solutions = document.add_table(rows=len(self.table), cols=len(self.table[0]) * 2)
        solutions.style = table_style

        for i in range(len(solutions.rows)):
            cells = solutions.rows[i].cells
            for j in range(int(len(cells) / 2)):
                cells[j * 2].text = f"{self.table[i][j].equation}"
                cells[j * 2 + 1].text = f"{self.table[i][j].solution}"

        document.save("math_problem_creator.docx")
